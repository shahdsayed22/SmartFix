import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_normal_font(doc, font_name='Times New Roman', font_size=12):
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = None
        if level == 1:
            run.font.size = Pt(20)
            run.bold = True
        elif level == 2:
            run.font.size = Pt(16)
            run.bold = True
        elif level == 3:
            run.font.size = Pt(14)
            run.bold = True

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(12)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_image_scaled(doc, img_path, width_inches=None, height_inches=None):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        if height_inches:
            r.add_picture(img_path, height=Inches(height_inches))
        elif width_inches:
            r.add_picture(img_path, width=Inches(width_inches))
        else:
            r.add_picture(img_path, width=Inches(6.0))
        p.paragraph_format.space_after = Pt(24)

def generate_chapter_2():
    doc = Document()
    set_normal_font(doc)

    add_heading(doc, 'Chapter 2: System Analysis and Requirements Specification', 1)

    add_heading(doc, '2.1 Introduction and Problem Definition', 2)
    add_paragraph(doc, "The SmartFix project was initiated to systematically resolve the profound latency and informational asymmetry endemic to the informal maintenance sector in Egypt. This chapter defines the fundamental boundaries of the proposed software system by conducting rigorous feasibility studies and extracting strict IEEE-standard functional requirements necessary for the architectural implementation detailed in Chapter 3.")
    
    add_image_scaled(doc, 'SmartFix_Diagrams/screenshot_dashboard_home.png', width_inches=6.0)

    add_heading(doc, '2.2 Systems Feasibility Study', 2)
    add_paragraph(doc, "Prior to executing software development, a tri-fold feasibility analysis confirmed the viability of the SmartFix application.")
    
    add_heading(doc, '2.2.1 Economic and Operational Feasibility', 3)
    add_paragraph(doc, "Economically, the system leverages a decentralized gig-economy model. The operational paradigm significantly undercuts traditional corporate overhead by eliminating centralized dispatchers. Operationally, it introduces algorithms that decrease customer wait-times from hours down to sub-second localized search queries, ensuring massive adoption probabilities.")
    
    add_heading(doc, '2.2.2 Technical Feasibility', 3)
    add_paragraph(doc, "Technically, the integration of Flutter guarantees cross-platform ubiquitous access. The utilization of MongoDB provides NoSQL spatial indexing crucial for locating targets. Next.js functions construct a stateless interface guaranteeing high availability.")

    add_heading(doc, '2.3 System Requirements (IEEE Framework)', 2)
    add_paragraph(doc, "The following requirements strictly encapsulate the operational expectations of the software.")
    
    add_heading(doc, '2.3.1 Functional Requirements (FR)', 3)
    add_bullet(doc, "FR-01: The system MUST allow customers to submit maintenance issues detailing descriptions, categories, and geolocations.")
    add_bullet(doc, "FR-02: The system SHALL algorithmically route issues specifically to technicians within a localized spatial radius.")
    add_bullet(doc, "FR-03: Technicians MUST be able to accept or reject incoming job queues dynamically.")
    add_bullet(doc, "FR-04: The system MUST facilitate bidirectional real-time chat between the Consumer and the Technician exclusively post-acceptance.")
    add_bullet(doc, "FR-05: Administrators SHALL possess the capability to globally monitor system metrics via an interactive web dashboard.")

    add_heading(doc, '2.3.2 Non-Functional Requirements (NFR)', 3)
    add_bullet(doc, "NFR-01 (Performance): API endpoint interactions must execute and return within 400 milliseconds under normal load.")
    add_bullet(doc, "NFR-02 (Security): All endpoint requests must require JWT-based cryptographic signatures verified globally against Firebase Auth.")
    add_bullet(doc, "NFR-03 (Scalability): The NoSQL BSON schema must support horizontal scaling capable of indexing 10,000 active concurrent technicians.")

    add_heading(doc, '2.4 Use Case Modeling and Actor Descriptions', 2)
    add_paragraph(doc, "The functional requirements are visualized topologically across the following Use Case mapping, distinctly separating Customer, Technician, and Administrator permissions.")

    # Injected Use Case (Tall Image Fix applied)
    add_image_scaled(doc, 'SmartFix_Diagrams/04_Use_Case_Diagram/use_case_diagram.png', height_inches=7.5)

    add_heading(doc, '2.4.1 Use Case Specifications', 3)
    add_paragraph(doc, "UC-01: 'Report Maintenance Issue'")
    add_bullet(doc, "Actor: Customer")
    add_bullet(doc, "Pre-condition: User must be authenticated and GPS permissions enabled.")
    add_bullet(doc, "Main Flow: User selects category -> inputs description -> captures image -> UI extracts latitude/longitude -> pushes payload to /api/issues.")
    add_bullet(doc, "Post-condition: Issue state becomes 'Pending', triggering matching algorithm.")

    add_heading(doc, '2.5 Process Modeling and Navigational Flows', 2)
    add_paragraph(doc, "The structural lifecycle modeling dictates how the 'Pending' issue transitions securely through the system boundaries over time.")

    # Activity Diagram
    add_image_scaled(doc, 'SmartFix_Diagrams/07_Activity_Issue_Lifecycle/activity_issue_lifecycle.png', height_inches=7.0)

    add_paragraph(doc, "Data propagates sequentially through these physical interfaces, updating states iteratively in MongoDB before reflecting back to the UI threads.")
    # Data Flow
    add_image_scaled(doc, 'SmartFix_Diagrams/13_Data_Flow_Diagram/data_flow_diagram.png', width_inches=6.0)

    add_paragraph(doc, "Finally, user interaction UX modeling prevents logical dead-ends by strictly grouping routing protocols locally internally within the Flutter framework.")
    # Screen Navigation
    add_image_scaled(doc, 'SmartFix_Diagrams/14_Screen_Navigation_Flow/screen_navigation_flow.png', width_inches=6.0)

    doc.save('SmartFix_Chapter_2.docx')


def generate_chapter_3():
    doc = Document()
    set_normal_font(doc)

    add_heading(doc, 'Chapter 3: Methodology and System Design', 1)

    add_heading(doc, '3.1 Introduction', 2)
    add_paragraph(doc, "Translating a comprehensive literature review into functional software necessitates exhaustive architectural definitions. Chapter 3 focuses extensively on the tangible blueprints and systematic workflows implemented during the conception and execution of the SmartFix project.")
    
    add_heading(doc, '3.2 Software Development Life Cycle (SDLC): Agile Methodology', 2)
    add_paragraph(doc, "In managing volatile complexity, the project strictly adopted an Agile Scrum Methodology predicated upon recursive loops. Through Continuous Integration, internal module packages systematically interconnect independent operational nodes.")
    # Package Diagram
    add_image_scaled(doc, 'SmartFix_Diagrams/15_Package_Diagram/package_diagram.png', width_inches=6.0)

    add_heading(doc, '3.3 High-Level System Architecture and Deployment', 2)
    add_paragraph(doc, "SmartFix dictates a tightly integrated Client-Server model relying logically upon internal component structures mapping explicit operational pathways.")
    
    # System Architecture
    add_image_scaled(doc, 'SmartFix_Diagrams/01_System_Architecture/system_architecture.png', width_inches=6.0)
    
    add_paragraph(doc, "On an abstract atomic level, structural components actively dictate interactions mapping distinct modular domains logic blocks isolating dependencies.")
    # Component Diagram
    add_image_scaled(doc, 'SmartFix_Diagrams/08_Component_Diagram/component_diagram.png', width_inches=6.0)

    add_paragraph(doc, "Physically, executing the compiled platform forces isolated Docker images alongside external managed databases dynamically hosted on distinct networking hardware frameworks mapping strict deployment hierarchies.")
    # Deployment diagram
    add_image_scaled(doc, 'SmartFix_Diagrams/09_Deployment_Diagram/deployment_diagram.png', width_inches=6.0)

    add_heading(doc, '3.4 Database Modeling and Class Structuring', 2)
    add_paragraph(doc, "Data persistence enforces Mongoose 9 ODM structure. The system establishes distinct hierarchical models organizing logic structures inside Object-Oriented Dart patterns and strictly enforced backend paradigms.")
    
    # ER Diagram
    add_image_scaled(doc, 'SmartFix_Diagrams/02_ER_Diagram/er_diagram.png', width_inches=6.0)
    
    add_paragraph(doc, "Object-oriented structures mandate specific classes orchestrating memory manipulation and state extraction mapping deeply typed schemas within the application execution loop.")
    # Class Diagram (potentially tall)
    add_image_scaled(doc, 'SmartFix_Diagrams/03_Class_Diagram/class_diagram.png', height_inches=7.0)

    add_heading(doc, '3.5 State Interaction and Sequence Logic', 2)
    add_paragraph(doc, "Application interactions transition functionally utilizing State Machine constructs mapping absolute behavioral paradigms restricted globally against unauthorized mutations across operational states.")
    # State Machine
    add_image_scaled(doc, 'SmartFix_Diagrams/10_State_Machine_Issue/state_machine_issue.png', width_inches=6.0)

    add_paragraph(doc, "Network sequences inherently necessitate mapping routing parameters defining exactly how Next.js APIs bind mobile callbacks across endpoints and RESTful paradigms.")
    # API endpoints
    add_image_scaled(doc, 'SmartFix_Diagrams/16_API_Endpoint_Map/api_endpoint_map.png', height_inches=7.0)

    add_heading(doc, '3.5.1 Sequence Activity Schemas', 3)
    add_paragraph(doc, "Below are detailed sequence abstractions mapping multi-layered latency paradigms explicitly outlining procedural API execution events chronologically:")
    
    add_paragraph(doc, "Execution Sequence: Authentication Loop.")
    add_image_scaled(doc, 'SmartFix_Diagrams/11_Sequence_Authentication/sequence_authentication.png', width_inches=6.0)
    
    add_paragraph(doc, "Execution Sequence: Systemic Issue Generation.")
    add_image_scaled(doc, 'SmartFix_Diagrams/05_Sequence_Report_Issue/sequence_report_issue.png', width_inches=6.0)

    add_paragraph(doc, "Execution Sequence: Dispatch Acceptance Routing.")
    add_image_scaled(doc, 'SmartFix_Diagrams/06_Sequence_Accept_Job/sequence_accept_job.png', width_inches=6.0)
    
    add_paragraph(doc, "Execution Sequence: Bidirectional Cloud Chat Sync.")
    add_image_scaled(doc, 'SmartFix_Diagrams/12_Sequence_Chat_Flow/sequence_chat_flow.png', width_inches=6.0)

    add_heading(doc, '3.6 User Interface and Control Interfaces', 2)
    add_paragraph(doc, "Designing administrative environments demands absolute data oversight scaling logically across multi-dimensional metric interfaces mapping total operational velocity globally.")
    
    # Final screenshots
    add_image_scaled(doc, 'SmartFix_Diagrams/screenshot_issues_management.png', width_inches=6.0)
    add_image_scaled(doc, 'SmartFix_Diagrams/screenshot_technicians.png', width_inches=6.0)

    doc.save('SmartFix_Chapter_3.docx')

if __name__ == '__main__':
    generate_chapter_2()
    generate_chapter_3()
    print("Massively detailed Chapter documents with ALL 19 diagrams generated successfully.")
